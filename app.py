
from flask import jsonify, send_file
from flask import Flask, render_template, request
from docx.oxml import OxmlElement
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import os
import json
from concurrent.futures import ThreadPoolExecutor
# 初始化线程池执行器
executor = ThreadPoolExecutor(max_workers=4)

def add_bookmark(paragraph, bookmark_name):
    """在段落中添加书签"""
    bookmark_id = str(uuid.uuid4())[:8]  # 生成较短的唯一 ID
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    bookmark_end.set(qn('w:name'), bookmark_name)

    paragraph._p.append( bookmark_start)
    paragraph._p.append( bookmark_end)

def add_hyperlink_to_bookmark(run,paragraph, bookmark_name, link_text, insert_position):
    """在段落中添加超链接到书签"""
    # 创建超链接
    hyperlink = OxmlElement('w:hyperlink')
    rId = str(uuid.uuid4())[:8]  # 生成较短的唯一 ID
    hyperlink.set(qn('r:id'), rId)  # 设置超链接的 ID
    hyperlink.set(qn('w:anchor'), bookmark_name)  # 指定书签名称

    # 将超链接添加到段落
    if insert_position is not None:
        if insert_position==0:
            new_run = paragraph.add_run(link_text)  # 使用 paragraph.add_run 创建运行
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            new_run.font.superscript = False
            hyperlink.append(new_run._r)
            paragraph._p.insert(insert_position, hyperlink)
        else:
            new_run = paragraph.add_run(link_text)  # 使用 paragraph.add_run 创建运行
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            new_run.font.superscript = True  # 设置运行文本为上标
            hyperlink.append(new_run._r)
            paragraph._p.insert(insert_position, hyperlink)
    else:
        new_run = paragraph.add_run(link_text)  # 使用 paragraph.add_run 创建运行
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.superscript = True  # 设置运行文本为上标
        hyperlink.append(new_run._r)
        paragraph._p.append(hyperlink)

def adjust_references(file_path, output_path,kuohao, reference_range):
    """处理文档，插入书签和超链接"""
    doc = Document(file_path)
    # 为每个引用创建书签
    reference_bookmarks = {}
    for reference in reference_range:
        if kuohao=='【】':
            reference_text = f"【{reference}】"
        elif kuohao=='[]':
            reference_text = f"[{reference}]"
        else:
            raise("请选择合适括号")

        bookmark_name = f"Reference{reference}"
        reference_bookmarks[reference_text] = bookmark_name

    # 查找每个引用所在的段落并添加书签
    for reference, bookmark_name in reference_bookmarks.items():
        last_reference_paragraph = None
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if reference in run.text:
                    last_reference_paragraph = paragraph
        if last_reference_paragraph:
            add_bookmark(last_reference_paragraph, bookmark_name)

    # 遍历文档中的段落，插入超链接

    All_add_sum={} #添加[1]上标的全段额外加一
    for reference, bookmark_name in reference_bookmarks.items():
        for paragraph in doc.paragraphs:
            sum=0
            if str(paragraph) not in All_add_sum:
                sum +=0
            else:
                sum +=All_add_sum[str(paragraph)]
               # print(All_add_sum[str(paragraph)])
            index = 0
            for run in paragraph.runs:
                index +=1
            if index==0:
                continue
            for run in paragraph.runs:
                sum+=1
                if reference in run.text:
                    if str(paragraph) not in All_add_sum:
                        All_add_sum[str(paragraph)]=1
                    else:
                        All_add_sum[str(paragraph)]=All_add_sum[str(paragraph)]+1
                   #     print('reference',reference)
                  #      print(All_add_sum[str(paragraph)])

                    if index==1:
                   #     print('end')
                        run.text = run.text.replace(reference,'')
                        add_hyperlink_to_bookmark(run,paragraph, bookmark_name, reference, 0)
                    else:
                        index -=1
                        run.text = run.text.replace(reference,'')
                        add_hyperlink_to_bookmark(run,paragraph, bookmark_name, reference,sum)
                        sum += 1  # 因为多添加了[i]

    # 保存调整后的文档
    doc.save(output_path)
    print(f"文档已保存到 {output_path}")

def adjust_docx_format(file_path, output_path,
                       list_head1, list_head2, list_head3, list_head4,
                       font_name='Arial', font_size=12,
                       heading1_font='黑体', heading1_size=10.5, heading1_bold=False,
                       heading2_font='Arial', heading2_size=14, heading2_bold=False,
                       heading3_font='Arial', heading3_size=12, heading3_bold=False,
                       heading4_font='Arial', heading4_size=11, heading4_bold=False):
    # 打开文档
    doc = Document(file_path)

    def split_paragraphs_simple(doc):
        """确保完全保持原始顺序的段落拆分"""
        paragraphs = list(doc.paragraphs)  # 创建段落副本

        for paragraph in paragraphs:
            # 检查是否需要处理（有换行或图片）
            if '\n' in paragraph.text or any('pic:pic' in run._element.xml for run in paragraph.runs):
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                original_runs = list(paragraph.runs)  # 保留原始run顺序

                # 先移除原段落
                parent.remove(paragraph._element)

                current_text = []
                current_para = None

                for run in original_runs:
                    # 处理图片run
                    if 'pic:pic' in run._element.xml:
                        # 如果有等待处理的文本，先创建段落
                        if current_text:
                            current_para = doc.add_paragraph(''.join(current_text))
                            parent.insert(index, current_para._element)
                            index += 1
                            current_text = []

                        # 创建独立的图片段落
                        pic_para = doc.add_paragraph()
                        pic_run = pic_para.add_run()
                        for child in run._element:
                            pic_run._element.append(child)
                        parent.insert(index, pic_para._element)
                        index += 1

                    # 处理文本run
                    elif run.text:
                        # 按换行符拆分文本
                        lines = run.text.split('\n')
                        for i, line in enumerate(lines):
                            if line.strip():
                                current_text.append(line)
                            if i < len(lines) - 1:  # 不是最后一行
                                if current_text:  # 如果有等待处理的文本
                                    current_para = doc.add_paragraph(''.join(current_text))
                                    parent.insert(index, current_para._element)
                                    index += 1
                                    current_text = []

                # 处理最后剩余的文本
                if current_text:
                    current_para = doc.add_paragraph(''.join(current_text))
                    parent.insert(index, current_para._element)



    # 定义递归处理文本的函数
    def process_paragraph(paragraph):
        text = paragraph.text.strip()
        if not text:
            return

        text_to_match =(text.split('\n')[0] if '\n' in text else text)[:8]

        if any(text_to_match.startswith(head) for head in list_head1):
            paragraph.style = doc.styles['Heading 1']
            for run in paragraph.runs:
                run.font.name = heading1_font
                run.font.size = Pt(heading1_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading1_font)
                run.bold = heading1_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head2):
            paragraph.style = doc.styles['Heading 2']
            for run in paragraph.runs:
                run.font.name = heading2_font
                run.font.size = Pt(heading2_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading2_font)
                run.bold = heading2_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head3):
            paragraph.style = doc.styles['Heading 3']
            for run in paragraph.runs:
                run.font.name = heading3_font
                run.font.size = Pt(heading3_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading3_font)
                run.bold = heading3_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head4):
            paragraph.style = doc.styles['Heading 4']
            for run in paragraph.runs:
                run.font.name = heading4_font
                run.font.size = Pt(heading4_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading4_font)
                run.bold = heading4_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        else:
            paragraph.style = doc.styles['Normal']
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False

    split_paragraphs_simple(doc)
    # 遍历文档中的段落并处理
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # 保存调整后的文档
    doc.save(output_path)
    print(f"文档已保存到 {output_path}")


def cross_reference(file_path,output_path,kuohao='[]',max=30):
    if kuohao =='无':
        doc = Document(file_path)
        doc.save(output_path)
        print(f"文档已保存到 {output_path}")
        return
    adjust_references(file_path, output_path,kuohao,reference_range=range(1, max+1))


#信息初始化############################################
app = Flask(__name__,static_folder='static')
# 确保上传目录存在
UPLOAD_FOLDER ='uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
print('app已启动')
directory_path = "uploads"  # 替换为你的目标目录路径
max_size = 1 * 1024 * 1024 *1024 *1024  # 1024GB，单位为字节

##信息转化###########################################
FONT_SIZE = {
        "初号": 42,
        "小初号": 36,
        "一号": 26,
        "小一号": 24,
        "二号": 22,
        "小二号": 18,
        "三号": 16,
        "小三号": 15,
        "四号": 14,
        "小四号": 12,
        "五号": 10.5,
        "小五号": 9,
        "六号": 7.5,
        "小六号": 6.5,
        "七号": 5.5,
        "八号": 5
    }
BOLD = {
        "是": True,
        "否": False
    }
def generate_list_heads(param_str):

    list0 = [
            ["$/**//**//***/$"],  # 0
            ['一、','二、','三、','四、','五、','六、','七、','八、','九、','十、'],#1
            ['（一）','（二）',"（三）","（四）","（五）","（六）","（七）","（八）","（九）","（十）"],#2
            ['1.','2.','3.','4.','5.','6.','7.','8.','9.','10.'],#3
            ['（1）','（2）','（3）','（4）','（5）','（6）','（7）','（8）','（9）','（10）'],#4
            ['1.1','1.2','1.3','1.4','1.5','1.6','1.7','1.8',
           '2.1','2.2','2.3','2.4','2.5','2.6','2.7','2.8',
           '3.1','3.2','3.3','3.4','3.5','3.6','3.7','3.8',
           '4.1','4.2','4.3','4.4','4.5','4.6','4.7','4.8',
           '5.1','5.2','5.3','5.4','5.5','5.6','5.7','5.8',
           '6.1','6.2','6.3','6.4','6.5','6.6','6.7','6.8',
           '7.1','7.2','7.3','7.4','7.5','7.6','7.7','7.8',
           '8.1','8.2','8.3','8.4','8.5','8.6','8.7','8.8'],  # 5
            ['1.1.1','1.1.2','1.1.3','1.1.4','1.2.1','1.2.2','1.2.3','1.2.4','1.3.1','1.3.2','1.3.3',
           '1.3.4','1.4.1','1.4.2','1.4.3','1.4.4',
           '2.1.1','2.1.2','2.1.3','2.1.4','2.2.1','2.2.2','2.2.3','2.2.4','2.3.1','2.3.2','2.3.3',
           '2.3.4','2.4.1','2.4.2','2.4.3','2.4.4',
           '3.1.1','3.1.2','3.1.3','3.1.4','3.2.1','3.2.2','3.2.3','3.2.4','3.3.1','3.3.2','3.3.3',
           '3.3.4','3.4.1','3.4.2','3.4.3','3.4.4',
           '4.1.1','4.1.2','4.1.3','4.1.4','4.2.1','4.2.2','4.2.3','4.2.4','4.3.1','4.3.2','4.3.3',
           '4.3.4','4.4.1','4.4.2','4.4.3','4.4.4','4.4.5','4.4.6','4.4.7','4.4.8','4.4.9','4.4.10','4.4.11',
           '4.5.1','4.5.2','4.5.3','4.5.4','4.5.5','4.5.6','4.5.7','4.5.8','4.5.9','4.5.10','4.5.11'],  # 6
            ['1.1.1.1','1.1.1.2','1.1.1.3','1.1.1.4','1.1.2.1','1.1.2.2','1.1.2.3','1.1.2.4','1.1.3.1',
           '1.1.3.2','1.1.3.3','1.1.3.4','1.1.4.1','1.1.4.2','1.1.4.3','1.1.4.4',
           '1.2.1.1','1.2.1.2','1.2.1.3','1.2.1.4','1.2.2.1','1.2.2.2','1.2.2.3','1.2.2.4','1.2.3.1',
           '1.2.3.2','1.2.3.3','1.2.3.4','1.2.4.1','1.2.4.2','1.2.4.3','1.2.4.4',
           '1.3.1.1','1.3.1.2','1.3.1.3','1.3.1.4','1.3.2.1','1.3.2.2','1.3.2.3','1.3.2.4','1.3.3.1',
           '1.3.3.2','1.3.3.3','1.3.3.4','1.3.4.1','1.3.4.2','1.3.4.3','1.3.4.4',
           '1.4.1.1','1.4.1.2','1.4.1.3','1.4.1.4','1.4.2.1','1.4.2.2','1.4.2.3','1.4.2.4','1.4.3.1',
            '1.4.3.2','1.4.3.3','1.4.3.4','1.4.4.1','1.4.4.2','1.4.4.3','1.4.4.4',
            '2.1.1.1','2.1.1.2','2.1.1.3','2.1.1.4','2.1.2.1','2.1.2.2','2.1.2.3','2.1.2.4','2.1.3.1',
            '2.1.3.2','2.1.3.3','2.1.3.4','2.1.4.1','2.1.4.2','2.1.4.3','2.1.4.4',
            '2.2.1.1','2.2.1.2','2.2.1.3','2.2.1.4','2.2.2.1','2.2.2.2','2.2.2.3','2.2.2.4','2.2.3.1',
            '2.2.3.2','2.2.3.3','2.2.3.4','2.2.4.1','2.2.4.2','2.2.4.3','2.2.4.4',
            '2.3.1.1','2.3.1.2','2.3.1.3','2.3.1.4','2.3.2.1','2.3.2.2','2.3.2.3','2.3.2.4','2.3.3.1',
            '2.3.3.2','2.3.3.3','2.3.3.4','2.3.4.1','2.3.4.2','2.3.4.3','2.3.4.4',
            '2.4.1.1','2.4.1.2','2.4.1.3','2.4.1.4','2.4.2.1','2.4.2.2','2.4.2.3','2.4.2.4','2.4.3.1',
            '2.4.3.2','2.4.3.3','2.4.3.4','2.4.4.1','2.4.4.2','2.4.4.3','2.4.4.4',
            '3.1.1.1','3.1.1.2','3.1.1.3','3.1.1.4','3.1.2.1','3.1.2.2','3.1.2.3','3.1.2.4','3.1.3.1',
            '3.1.3.2','3.1.3.3','3.1.3.4','3.1.4.1','3.1.4.2','3.1.4.3','3.1.4.4',
            '3.2.1.1','3.2.1.2','3.2.1.3','3.2.1.4','3.2.2.1','3.2.2.2','3.2.2.3','3.2.2.4','3.2.3.1',
            '3.2.3.2','3.2.3.3','3.2.3.4','3.2.4.1','3.2.4.2','3.2.4.3','3.2.4.4',
            '3.3.1.1','3.3.1.2','3.3.1.3','3.3.1.4','3.3.2.1','3.3.2.2','3.3.2.3','3.3.2.4','3.3.3.1',
            '3.3.3.2','3.3.3.3','3.3.3.4','3.3.4.1','3.3.4.2','3.3.4.3','3.3.4.4',
            '3.4.1.1','3.4.1.2','3.4.1.3','3.4.1.4','3.4.2.1','3.4.2.2','3.4.2.3','3.4.2.4','3.4.3.1',
            '3.4.3.2','3.4.3.3','3.4.3.4','3.4.4.1','3.4.4.2','3.4.4.3','3.4.4.4',
            '4.1.1.1','4.1.1.2','4.1.1.3','4.1.1.4','4.1.2.1','4.1.2.2','4.1.2.3','4.1.2.4','4.1.3.1',
            '4.1.3.2','4.1.3.3','4.1.3.4','4.1.4.1','4.1.4.2','4.1.4.3','4.1.4.4',
            '4.2.1.1','4.2.1.2','4.2.1.3','4.2.1.4','4.2.2.1','4.2.2.2','4.2.2.3','4.2.2.4','4.2.3.1',
            '4.2.3.2','4.2.3.3','4.2.3.4','4.2.4.1','4.2.4.2','4.2.4.3','4.2.4.4',
            '4.3.1.1','4.3.1.2','4.3.1.3','4.3.1.4','4.3.2.1','4.3.2.2','4.3.2.3','4.3.2.4','4.3.3.1',
            '4.3.3.2','4.3.3.3','4.3.3.4','4.3.4.1','4.3.4.2','4.3.4.3','4.3.4.4',
            '4.4.1.1','4.4.1.2','4.4.1.3','4.4.1.4','4.4.2.1','4.4.2.2','4.4.2.3','4.4.2.4','4.4.3.1',
            '4.4.3.2','4.4.3.3','4.4.3.4','4.4.4.1','4.4.4.2','4.4.4.3','4.4.4.4','4.5.1.1','4.5.1.2',
            '4.5.1.3','4.5.1.4','4.5.2.1','4.5.2.2','4.5.2.3','4.5.2.4','4.5.3.1','4.5.3.2','4.5.3.3',
            '4.5.3.4','4.5.4.1','4.5.4.2','4.5.4.3','4.5.4.4'],# 7

            ['I','II','III','IV','V','VI','VII' ,'VIII','IX','X'],  # 8
            ['A','B','C','D','E','F','G','H','I','J'],  # 9
            ['a','b','c','d','e','f','g','h','i','j'],  # 10
            ['(A)','(B)','(C)','(D)','(E)','(F)','(G)','(H)','(I)','(J)'],  # 11
            ['(a)','(b)','(c)','(d)','(e)','(f)','(g)','(h)','(i)','(j)'],  # 12
            ['1-1','1-2','1-3','1-4','1-5','1-6','1-7','1-8','1-9','1-10'],  # 13
            ['1-1-1','1-1-2','1-1-3','1-1-4','1-1-5','1-1-6','1-1-7','1-1-8','1-1-9','1-1-10'],  # 14
            ['1-1-1-1','1-1-1-2','1-1-1-3','1-1-1-4','1-1-1-5','1-1-1-6','1-1-1-7','1-1-1-8','1-1-1-9','1-1-1-10'],# 15
            ['1.','2.','3.','4.','5.','6.','7.','8.','9.','10.'],  # 16
            ['1、','2、','3、','4、','5、','6、','7、','8、','9、','10、'],  # 17
            ['(1)','(2)','(3)','(4)','(5)','(6)','(7)','(8)','(9)','(10)'],  # 18
            ['（1）','（2）','（3）','（4）','（5）','（6）','（7）','（8）','（9）','（10）'],  # 19
    ]
    ele=param_str.split(' ')[0]
    if ele=="无":
        return list0[0]
    for list_num in list0:
      if ele in list_num:
        return list_num
def generate_kuohao(param_str):
    if "无" in param_str:
        return "无"
    ele=param_str.split('')[0]
    if ele=="[1]":
            return "[]"
    return "【】"

@app.route('/process2', methods=['POST'])
def process2():
############################传递文件名称
  try:
    if 'document' not in request.files:
        return jsonify({"error": "没有文件部分"}), 400
    file = request.files['document']
    if file.filename =='':
        return jsonify({"error": "没有选择文件"}), 400
    # 获取文件扩展名
    file_extension = os.path.splitext(file.filename)[1].lower()
    file_ori_name=os.path.splitext(file.filename)[0].lower()
    # 检查文件扩展名是否为 .docx 或 .doc
    if file_extension not in ['.docx','.doc']:
        return jsonify({"error": "文件格式不正确，仅支持 .docx 或 .doc 格式"}), 400

    # 获取文件的原始文件名
    filename = file_ori_name+'-'+str(uuid.uuid4())+file_extension
    # 构造保存路径
    save_path = os.path.join(UPLOAD_FOLDER, filename)
    # 保存文件到指定路径
    file.save(save_path)
    file_path =save_path 
    file_name, file_extension = os.path.splitext(file_path)
    
    # 生成 temp_path
    temp_path = f"{file_name}temp{file_extension}"
    # 生成 output_path
    output_path = f"{file_name}out_put{file_extension}"

##########################################传递其余参数

    params = json.loads(request.form['params'])
    # 自定义标题内容列表
    list_head1 =generate_list_heads(params['title1_label'])
    list_head2 =generate_list_heads(params['title2_label'])
    list_head3 =generate_list_heads(params['title3_label'])
    list_head4 =generate_list_heads(params['title4_label'])

    # 调用函数
    future = executor.submit(adjust_docx_format,
        file_path, temp_path,
        list_head1, list_head2, list_head3, list_head4,
        font_name=params['body_font'], font_size=FONT_SIZE[params['body_size']],
        heading1_font=params['title1_font'], heading1_size=FONT_SIZE[params['title1_size']], heading1_bold=BOLD[params['title1_bold']],
        heading2_font=params['title2_font'], heading2_size=FONT_SIZE[params['title2_size']], heading2_bold=BOLD[params['title2_bold']],
        heading3_font=params['title3_font'], heading3_size=FONT_SIZE[params['title3_size']], heading3_bold=BOLD[params['title3_bold']],
        heading4_font=params['title4_font'], heading4_size=FONT_SIZE[params['title4_size']], heading4_bold=BOLD[params['title4_bold']]
    )
    future.result(timeout=10)
    cross_reference(temp_path,output_path,kuohao=generate_kuohao(params['citation_style']),max=30)
    #接口用法
    os.remove(temp_path)
    print(f"已删除: {temp_path}")
    return send_file(
            output_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='formatted_document.docx'
    )

  except Exception as e:
        return jsonify({"服务器报错": str(e)}), 500

@app.route('/')
def index():
    return render_template('index.html')

if __name__ =='__main__':
   # app.run(host='0.0.0.0', port=5000, debug=True, threads=4)
    app.run(host='127.0.0.1', port=5000, debug=True)

