from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import uuid
from concurrent.futures import ThreadPoolExecutor

# 所有可选项定义

# 字号转换为磅值
FONT_SIZE_MAP = {
    "初号": 42, "小初号": 36, "一号": 26, "小一号": 24,
    "二号": 22, "小二号": 18, "三号": 16, "小三号": 15,
    "四号": 14, "小四号": 12, "五号": 10.5, "小五号": 9,
    "六号": 7.5, "小六号": 6.5, "七号": 5.5, "八号": 5
}

# 初始化线程池执行器
executor = ThreadPoolExecutor(max_workers=1)


def add_bookmark(paragraph, bookmark_name):
    """在段落中添加书签"""
    bookmark_id = str(uuid.uuid4())[:8]
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    bookmark_end.set(qn('w:name'), bookmark_name)

    paragraph._p.append(bookmark_start)
    paragraph._p.append(bookmark_end)


def add_hyperlink_to_bookmark(run, paragraph, bookmark_name, link_text, insert_position):
    """在段落中添加超链接到书签"""
    hyperlink = OxmlElement('w:hyperlink')
    rId = str(uuid.uuid4())[:8]
    hyperlink.set(qn('r:id'), rId)
    hyperlink.set(qn('w:anchor'), bookmark_name)

    if insert_position is not None:
        new_run = paragraph.add_run(link_text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.superscript = (insert_position != 0)
        hyperlink.append(new_run._r)
        paragraph._p.insert(insert_position, hyperlink)
    else:
        new_run = paragraph.add_run(link_text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.superscript = True
        hyperlink.append(new_run._r)
        paragraph._p.append(hyperlink)


def adjust_references(file_path, output_path, kuohao, reference_range):
    """处理文档，插入书签和超链接"""
    doc = Document(file_path)

    # 为每个引用创建书签
    reference_bookmarks = {}
    for reference in reference_range:
        reference_text = f"【{reference}】" if kuohao == '【】' else f"[{reference}]"
        bookmark_name = f"Reference{reference}"
        reference_bookmarks[reference_text] = bookmark_name

    # 查找并添加书签
    for reference, bookmark_name in reference_bookmarks.items():
        last_reference_paragraph = None
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if reference in run.text:
                    last_reference_paragraph = paragraph
        if last_reference_paragraph:
            add_bookmark(last_reference_paragraph, bookmark_name)

    # 添加超链接
    All_add_sum = {}
    for reference, bookmark_name in reference_bookmarks.items():
        for paragraph in doc.paragraphs:
            sum_pos = All_add_sum.get(str(paragraph), 0)
            run_count = len(paragraph.runs)

            if run_count == 0:
                continue

            for run in paragraph.runs:
                sum_pos += 1
                if reference in run.text:
                    All_add_sum[str(paragraph)] = All_add_sum.get(str(paragraph), 0) + 1
                    run.text = run.text.replace(reference, '')
                    insert_pos = 0 if run_count == 1 else sum_pos
                    add_hyperlink_to_bookmark(run, paragraph, bookmark_name, reference, insert_pos)
                    sum_pos += 1  # 因为多添加了[i]

    doc.save(output_path)
    print(f"文档已保存到 {output_path}")


def adjust_docx_format(file_path, output_path,
                       list_head1, list_head2, list_head3, list_head4,
                       font_name='Arial', font_size=12,
                       heading1_font='黑体', heading1_size=10.5, heading1_bold=False,
                       heading2_font='Arial', heading2_size=14, heading2_bold=False,
                       heading3_font='Arial', heading3_size=12, heading3_bold=False,
                       heading4_font='Arial', heading4_size=11, heading4_bold=False):
    doc = Document(file_path)

    def process_paragraph(paragraph):
        text = paragraph.text.strip()
        if not text:
            return

        text_to_match =(text.split('\n')[0] if '\n' in text else text)[:5]

        if any(text_to_match.startswith(head) for head in list_head1):
            paragraph.style = doc.styles['Heading 1']
            for run in paragraph.runs:
                run.font.name = heading1_font
                run.font.size = Pt(heading1_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading1_font)
                run.bold = heading1_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head2):
            paragraph.style = doc.styles['Heading 2']
            for run in paragraph.runs:
                run.font.name = heading2_font
                run.font.size = Pt(heading2_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading2_font)
                run.bold = heading2_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head3):
            paragraph.style = doc.styles['Heading 3']
            for run in paragraph.runs:
                run.font.name = heading3_font
                run.font.size = Pt(heading3_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading3_font)
                run.bold = heading3_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        elif any(text_to_match.startswith(head) for head in list_head4):
            paragraph.style = doc.styles['Heading 4']
            for run in paragraph.runs:
                run.font.name = heading4_font
                run.font.size = Pt(heading4_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading4_font)
                run.bold = heading4_bold
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False
        else:
            paragraph.style = doc.styles['Normal']
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.italic = False

    def split_paragraphs_simple(doc):
        """确保完全保持原始顺序的段落拆分"""
        paragraphs = list(doc.paragraphs)  # 创建段落副本

        for paragraph in paragraphs:
            # 检查是否需要处理（有换行或图片）
            if '\n' in paragraph.text or any('pic:pic' in run._element.xml for run in paragraph.runs):
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                original_runs = list(paragraph.runs)  # 保留原始run顺序

                # 先移除原段落
                parent.remove(paragraph._element)

                current_text = []
                current_para = None

                for run in original_runs:
                    # 处理图片run
                    if 'pic:pic' in run._element.xml:
                        # 如果有等待处理的文本，先创建段落
                        if current_text:
                            current_para = doc.add_paragraph(''.join(current_text))
                            parent.insert(index, current_para._element)
                            index += 1
                            current_text = []

                        # 创建独立的图片段落
                        pic_para = doc.add_paragraph()
                        pic_run = pic_para.add_run()
                        for child in run._element:
                            pic_run._element.append(child)
                        parent.insert(index, pic_para._element)
                        index += 1

                    # 处理文本run
                    elif run.text:
                        # 按换行符拆分文本
                        lines = run.text.split('\n')
                        for i, line in enumerate(lines):
                            if line.strip():
                                current_text.append(line)
                            if i < len(lines) - 1:  # 不是最后一行
                                if current_text:  # 如果有等待处理的文本
                                    current_para = doc.add_paragraph(''.join(current_text))
                                    parent.insert(index, current_para._element)
                                    index += 1
                                    current_text = []

                # 处理最后剩余的文本
                if current_text:
                    current_para = doc.add_paragraph(''.join(current_text))
                    parent.insert(index, current_para._element)


    split_paragraphs_simple(doc)
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    doc.save(output_path)
    print(f"文档已保存到 {output_path}")


def cross_reference(file_path, output_path, kuohao='[]', max_ref=30):
    adjust_references(file_path, output_path, kuohao, range(1, max_ref + 1))


def process_document_locally(input_file, output_file, params):
    """本地处理文档的主函数"""
    try:
        # 生成临时文件路径
        file_name, file_extension = os.path.splitext(input_file)
        temp_path = f"{file_name}temp{file_extension}"

        # 获取标题标识列表
        list_head1 = OPTIONS['title_label'][params['title1_label_idx']]
        list_head2 = OPTIONS['title_label'][params['title2_label_idx']]
        list_head3 = OPTIONS['title_label'][params['title3_label_idx']]
        list_head4 = OPTIONS['title_label'][params['title4_label_idx']]

        # 调整文档格式
        future = executor.submit(
            adjust_docx_format,
            input_file, temp_path,
            list_head1, list_head2, list_head3, list_head4,
            font_name=OPTIONS['body_font'][params['body_font_idx']],
            font_size=FONT_SIZE_MAP[OPTIONS['body_size'][params['body_size_idx']]],
            heading1_font=OPTIONS['title_font'][params['title1_font_idx']],
            heading1_size=FONT_SIZE_MAP[OPTIONS['title_size'][params['title1_size_idx']]],
            heading1_bold=(params['title1_bold_idx'] == 1),
            heading2_font=OPTIONS['title_font'][params['title2_font_idx']],
            heading2_size=FONT_SIZE_MAP[OPTIONS['title_size'][params['title2_size_idx']]],
            heading2_bold=(params['title2_bold_idx'] == 1),
            heading3_font=OPTIONS['title_font'][params['title3_font_idx']],
            heading3_size=FONT_SIZE_MAP[OPTIONS['title_size'][params['title3_size_idx']]],
            heading3_bold=(params['title3_bold_idx'] == 1),
            heading4_font=OPTIONS['title_font'][params['title4_font_idx']],
            heading4_size=FONT_SIZE_MAP[OPTIONS['title_size'][params['title4_size_idx']]],
            heading4_bold=(params['title4_bold_idx'] == 1)
        )
        future.result(timeout=60)

        # 处理交叉引用
        cross_reference(
            temp_path, output_file,
            kuohao=OPTIONS['citation_style'][params['citation_style_idx']],
            max_ref=30
        )

        # 删除临时文件
        os.remove(temp_path)
        print(f"已删除临时文件: {temp_path}")

        return True

    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        return False
    finally:
        pass




# 所有可选项定义
OPTIONS = {
    # 正文字体选项
    'body_font': [
        '黑体',# 0
        '仿宋',# 1
        '微软雅黑',# 2
        '隶书', # 3
        '幼圆',# 4
        '宋体',# 5
        '楷体', # 6
        'Arial', # 7
        'Calibri'# 8
    ],

    # 正文字号选项
    'body_size': [
        '初号', # 0
        '小初号',#1
        '一号',#2
        '小一号', #3
        '二号',#4
        '小二号',#5
        '三号', #6
        '小三号', #7
        '四号',#8
        '小四号',#9
        '五号',#10
        '小五号',#11
        '六号', #12
        '小六号', #13
        '七号',#14
        '八号',#15
    ],

    # 标题标识选项
    'title_label': [
        ['无'],  # 0
        ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、'],  # 1
        ['（一）', '（二）', "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）"],  # 2
        ['1 .', '2 .', '3 .', '4 .', '5 .', '6 .', '7 .', '8 .', '9 .', '10 .'],  # 3
        ['（1）', '（2）', '（3）', '（4）', '（5）', '（6）', '（7）', '（8）', '（9）', '（10）'],  # 4
        ['1.1 ', '1.2 ', '1.3 ', '1.4 ', '1.5 ', '1.6 ', '1.7 ', '1.8 ',
         '2.1 ', '2.2 ', '2.3 ', '2.4 ', '2.5 ', '2.6 ', '2.7 ', '2.8 ',
         '3.1 ', '3.2 ', '3.3 ', '3.4 ', '3.5 ', '3.6 ', '3.7 ', '3.8 ',
         '4.1 ', '4.2 ', '4.3 ', '4.4 ', '4.5 ', '4.6 ', '4.7 ', '4.8 ',
         '5.1 ', '5.2 ', '5.3 ', '5.4 ', '5.5 ', '5.6 ', '5.7 ', '5.8 ',
         '6.1 ', '6.2 ', '6.3 ', '6.4 ', '6.5 ', '6.6 ', '6.7 ', '6.8 ',
         '7.1 ', '7.2 ', '7.3 ', '7.4 ', '7.5 ', '7.6 ', '7.7 ', '7.8 ',
         '8.1 ', '8.2 ', '8.3 ', '8.4 ', '8.5 ', '8.6 ', '8.7 ', '8.8 '],  # 5
        ['1.1.1 ', '1.1.2 ', '1.1.3 ', '1.1.4 ', '1.2.1 ', '1.2.2 ', '1.2.3 ', '1.2.4 ', '1.3.1 ', '1.3.2 ', '1.3.3 ',
         '1.3.4 ', '1.4.1 ', '1.4.2 ', '1.4.3 ', '1.4.4 ',
         '2.1.1 ', '2.1.2 ', '2.1.3 ', '2.1.4 ', '2.2.1 ', '2.2.2 ', '2.2.3 ', '2.2.4 ', '2.3.1 ', '2.3.2 ', '2.3.3 ',
         '2.3.4 ', '2.4.1 ', '2.4.2 ', '2.4.3 ', '2.4.4 ',
         '3.1.1 ', '3.1.2 ', '3.1.3 ', '3.1.4 ', '3.2.1 ', '3.2.2 ', '3.2.3 ', '3.2.4 ', '3.3.1 ', '3.3.2 ', '3.3.3 ',
         '3.3.4 ', '3.4.1 ', '3.4.2 ', '3.4.3 ', '3.4.4 ',
         '4.1.1 ', '4.1.2 ', '4.1.3 ', '4.1.4 ', '4.2.1 ', '4.2.2 ', '4.2.3 ', '4.2.4 ', '4.3.1 ', '4.3.2 ', '4.3.3 ',
         '4.3.4 ', '4.4.1 ', '4.4.2 ', '4.4.3 ', '4.4.4 '],  # 6
        ['1.1.1.1 ', '1.1.1.2 ', '1.1.1.3 ', '1.1.1.4 ', '1.1.2.1 ', '1.1.2.2 ', '1.1.2.3 ', '1.1.2.4 ', '1.1.3.1 ',
         '1.1.3.2 ', '1.1.3.3 ', '1.1.3.4 ', '1.1.4.1 ', '1.1.4.2 ', '1.1.4.3 ', '1.1.4.4 ',
         '1.2.1.1 ', '1.2.1.2 ', '1.2.1.3 ', '1.2.1.4 ', '1.2.2.1 ', '1.2.2.2 ', '1.2.2.3 ', '1.2.2.4 ', '1.2.3.1 ',
         '1.2.3.2 ', '1.2.3.3 ', '1.2.3.4 ', '1.2.4.1 ', '1.2.4.2 ', '1.2.4.3 ', '1.2.4.4 ',
         '1.3.1.1 ', '1.3.1.2 ', '1.3.1.3 ', '1.3.1.4 ', '1.3.2.1 ', '1.3.2.2 ', '1.3.2.3 ', '1.3.2.4 ', '1.3.3.1 ',
         '1.3.3.2 ', '1.3.3.3 ', '1.3.3.4 ', '1.3.4.1 ', '1.3.4.2 ', '1.3.4.3 ', '1.3.4.4 ',
         '1.4.1.1 ', '1.4.1.2 ', '1.4.1.3 ', '1.4.1.4 ', '1.4.2.1 ', '1.4.2.2 ', '1.4.2.3 ', '1.4.2.4 ', '1.4.3.1 ',
         '1.4.3.2 ', '1.4.3.3 ', '1.4.3.4 ', '1.4.4.1 ', '1.4.4.2 ', '1.4.4.3 ', '1.4.4.4 ',
         '2.1.1.1 ', '2.1.1.2 ', '2.1.1.3 ', '2.1.1.4 ', '2.1.2.1 ', '2.1.2.2 ', '2.1.2.3 ', '2.1.2.4 ', '2.1.3.1 ',
         '2.1.3.2 ', '2.1.3.3 ', '2.1.3.4 ', '2.1.4.1 ', '2.1.4.2 ', '2.1.4.3 ', '2.1.4.4 ',
         '2.2.1.1 ', '2.2.1.2 ', '2.2.1.3 ', '2.2.1.4 ', '2.2.2.1 ', '2.2.2.2 ', '2.2.2.3 ', '2.2.2.4 ', '2.2.3.1 ',
         '2.2.3.2 ', '2.2.3.3 ', '2.2.3.4 ', '2.2.4.1 ', '2.2.4.2 ', '2.2.4.3 ', '2.2.4.4 ',
         '2.3.1.1 ', '2.3.1.2 ', '2.3.1.3 ', '2.3.1.4 ', '2.3.2.1 ', '2.3.2.2 ', '2.3.2.3 ', '2.3.2.4 ', '2.3.3.1 ',
         '2.3.3.2 ', '2.3.3.3 ', '2.3.3.4 ', '2.3.4.1 ', '2.3.4.2 ', '2.3.4.3 ', '2.3.4.4 ',
         '2.4.1.1 ', '2.4.1.2 ', '2.4.1.3 ', '2.4.1.4 ', '2.4.2.1 ', '2.4.2.2 ', '2.4.2.3 ', '2.4.2.4 ', '2.4.3.1 ',
         '2.4.3.2 ', '2.4.3.3 ', '2.4.3.4 ', '2.4.4.1 ', '2.4.4.2 ', '2.4.4.3 ', '2.4.4.4 ',
         '3.1.1.1 ', '3.1.1.2 ', '3.1.1.3 ', '3.1.1.4 ', '3.1.2.1 ', '3.1.2.2 ', '3.1.2.3 ', '3.1.2.4 ', '3.1.3.1 ',
         '3.1.3.2 ', '3.1.3.3 ', '3.1.3.4 ', '3.1.4.1 ', '3.1.4.2 ', '3.1.4.3 ', '3.1.4.4 ',
         '3.2.1.1 ', '3.2.1.2 ', '3.2.1.3 ', '3.2.1.4 ', '3.2.2.1 ', '3.2.2.2 ', '3.2.2.3 ', '3.2.2.4 ', '3.2.3.1 ',
         '3.2.3.2 ', '3.2.3.3 ', '3.2.3.4 ', '3.2.4.1 ', '3.2.4.2 ', '3.2.4.3 ', '3.2.4.4 ',
         '3.3.1.1 ', '3.3.1.2 ', '3.3.1.3 ', '3.3.1.4 ', '3.3.2.1 ', '3.3.2.2 ', '3.3.2.3 ', '3.3.2.4 ', '3.3.3.1 ',
         '3.3.3.2 ', '3.3.3.3 ', '3.3.3.4 ', '3.3.4.1 ', '3.3.4.2 ', '3.3.4.3 ', '3.3.4.4 ',
         '3.4.1.1 ', '3.4.1.2 ', '3.4.1.3 ', '3.4.1.4 ', '3.4.2.1 ', '3.4.2.2 ', '3.4.2.3 ', '3.4.2.4 ', '3.4.3.1 ',
         '3.4.3.2 ', '3.4.3.3 ', '3.4.3.4 ', '3.4.4.1 ', '3.4.4.2 ', '3.4.4.3 ', '3.4.4.4 ',
         '4.1.1.1 ', '4.1.1.2 ', '4.1.1.3 ', '4.1.1.4 ', '4.1.2.1 ', '4.1.2.2 ', '4.1.2.3 ', '4.1.2.4 ', '4.1.3.1 ',
         '4.1.3.2 ', '4.1.3.3 ', '4.1.3.4 ', '4.1.4.1 ', '4.1.4.2 ', '4.1.4.3 ', '4.1.4.4 ',
         '4.2.1.1 ', '4.2.1.2 ', '4.2.1.3 ', '4.2.1.4 ', '4.2.2.1 ', '4.2.2.2 ', '4.2.2.3 ', '4.2.2.4 ', '4.2.3.1 ',
         '4.2.3.2 ', '4.2.3.3 ', '4.2.3.4 ', '4.2.4.1 ', '4.2.4.2 ', '4.2.4.3 ', '4.2.4.4 ',
         '4.3.1.1 ', '4.3.1.2 ', '4.3.1.3 ', '4.3.1.4 ', '4.3.2.1 ', '4.3.2.2 ', '4.3.2.3 ', '4.3.2.4 ', '4.3.3.1 ',
         '4.3.3.2 ', '4.3.3.3 ', '4.3.3.4 ', '4.3.4.1 ', '4.3.4.2 ', '4.3.4.3 ', '4.3.4.4 ',
         '4.4.1.1 ', '4.4.1.2 ', '4.4.1.3 ', '4.4.1.4 ', '4.4.2.1 ', '4.4.2.2 ', '4.4.2.3 ', '4.4.2.4 ', '4.4.3.1 ',
         '4.4.3.2 ', '4.4.3.3 ', '4.4.3.4 ', '4.4.4.1 ', '4.4.4.2 ', '4.4.4.3 ', '4.4.4.4 '],# 7

        ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X'],  # 8
        ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J'],  # 9
        ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j'],  # 10
        ['(A)', '(B)', '(C)', '(D)', '(E)', '(F)', '(G)', '(H)', '(I)', '(J)'],  # 11
        ['(a)', '(b)', '(c)', '(d)', '(e)', '(f)', '(g)', '(h)', '(i)', '(j)'],  # 12
        ['1-1', '1-2', '1-3', '1-4', '1-5', '1-6', '1-7', '1-8', '1-9', '1-10'],  # 13
        ['1-1-1', '1-1-2', '1-1-3', '1-1-4', '1-1-5', '1-1-6', '1-1-7', '1-1-8', '1-1-9', '1-1-10'],  # 14
        ['1-1-1-1', '1-1-1-2', '1-1-1-3', '1-1-1-4', '1-1-1-5', '1-1-1-6', '1-1-1-7', '1-1-1-8', '1-1-1-9', '1-1-1-10'],# 15
        ['1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.'],  # 16
        ['1、', '2、', '3、', '4、', '5、', '6、', '7、', '8、', '9、', '10、'],  # 17
        ['(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)', '(10)'],  # 18
        ['（1）', '（2）', '（3）', '（4）', '（5）', '（6）', '（7）', '（8）', '（9）', '（10）'],  # 19
    ],

    # 标题字体选项（与正文字体相同）
    'title_font': [
        '黑体', '仿宋', '微软雅黑', '隶书', '幼圆',
        '宋体', '楷体', 'Arial', 'Calibri'
    ],

    # 标题字号选项（与正文字号相同）
    'title_size': [
        '初号', '小初号', '一号', '小一号', '二号', '小二号',
        '三号', '小三号', '四号', '小四号', '五号', '小五号',
        '六号', '小六号', '七号', '八号'
    ],

    # 是否加粗选项
    'bold': ['否', '是'],

    # 引用标识样式选项
    'citation_style': ['[]', '【】']
}


if __name__ == "__main__":
    # 定义参数索引（对应OPTIONS中的选项顺序）
    params = {
        # 正文字体：选择第5个选项（宋体）
        'body_font_idx': 5,

        # 正文字号：选择第9个选项（小四号）
        'body_size_idx': 9,



        # 一级标题标识：'一、',
        'title1_label_idx': 1,

        # 一级标题字体：选择第0个选项（黑体）
        'title1_font_idx': 0,

        # 一级标题字号：四号）
        'title1_size_idx': 8,

        # 一级标题是否加粗：选择第1个选项（是）
        'title1_bold_idx': 1,






        # 二级标题标识：（一）
        'title2_label_idx': 2,

        # 二级标题字体：选择第0个选项（黑体）
        'title2_font_idx': 0,

        # 二级标题字号：选择第3个选项（小四号）
        'title2_size_idx': 9,

        # 二级标题是否加粗：选择第1个选项（是）
        'title2_bold_idx': 0,





        # 三级标题标识：1.
        'title3_label_idx': 16,

        # 三级标题字体：选择第5个选（黑体）
        'title3_font_idx': 0,

        # 三级标题字号：选择第9个选项（小四号）
        'title3_size_idx': 9,

        # 三级标题是否加粗：选择第0个选项（否）
        'title3_bold_idx': 0,





        # 四级标题标识：（1）
        'title4_label_idx': 19,

        # 四级标题字体：选择第5个选项（宋体）
        'title4_font_idx': 5,

        # 四级标题字号：选择第9个选项（小四号）
        'title4_size_idx': 9,

        # 四级标题是否加粗：选择第0个选项（否）
        'title4_bold_idx': 0,

        # 引用标识样式：选择第0个选项（[]）
        'citation_style_idx': 0
    }
    # 输入和输出文件路径
    input_file =r"D:\homework\jiwang_quic\1.docx" # 替换为你的输入文件路径
    output_file = r"D:\homework\jiwang_quic\2.docx" # 替换为你的输出文件路径

    # 处理文档
    success = process_document_locally(input_file, output_file, params)

    if success:
        print("文档处理完成！输出文件:", output_file)
    else:
        print("文档处理失败")